"""Parse PDF/DOCX into per-page text and extracted tables."""
from __future__ import annotations

import io
import re
import unicodedata
from dataclasses import dataclass, field

import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument

# Below this, a PDF page is probably a scan -> OCR fallback (Phase 4).
MIN_TEXT_CHARS = 50

# Control chars (except \n \r \t) leak in from PDF bullet/dingbat glyphs.
_CONTROL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_MULTI_BLANK = re.compile(r"\n{3,}")
_TRAILING_WS = re.compile(r"[ \t]+\n")


def clean_text(text: str) -> str:
    """Strip extraction artifacts that would otherwise be embedded and quoted back.

    Note: glyphs a PDF's font cmap maps incorrectly (e.g. an embedded subset font
    rendering the rupee sign as 'I') cannot be recovered here — the wrong
    codepoint is what the file actually contains. OCR (Phase 4) is the fallback
    for those documents.
    """
    if not text:
        return ""
    text = unicodedata.normalize("NFKC", text)
    text = _CONTROL.sub(" ", text)
    text = text.replace("�", " ")  # unmappable glyph
    text = text.replace(" ", " ")  # non-breaking space
    text = _TRAILING_WS.sub("\n", text)
    text = _MULTI_BLANK.sub("\n\n", text)
    return text.strip()


@dataclass
class Table:
    page: int
    rows: list[list[str | None]]

    def to_markdown(self) -> str:
        if not self.rows:
            return ""
        head, *body = self.rows
        cells = [clean_text(c) if c else "" for c in head]
        out = ["| " + " | ".join(cells) + " |", "| " + " | ".join("---" for _ in cells) + " |"]
        for r in body:
            out.append("| " + " | ".join(clean_text(c) if c else "" for c in r) + " |")
        return "\n".join(out)


@dataclass
class Page:
    number: int
    text: str = ""
    needs_ocr: bool = False


@dataclass
class Parsed:
    pages: list[Page] = field(default_factory=list)
    tables: list[Table] = field(default_factory=list)

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def has_tables(self) -> bool:
        return len(self.tables) > 0


def parse(data: bytes, filename: str) -> Parsed:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _parse_pdf(data)
    if name.endswith(".docx"):
        return _parse_docx(data)
    raise ValueError(f"Unsupported file type: {filename}. Upload a PDF or DOCX.")


def _parse_pdf(data: bytes) -> Parsed:
    out = Parsed()
    with fitz.open(stream=data, filetype="pdf") as doc:
        for i, page in enumerate(doc, start=1):
            text = clean_text(page.get_text() or "")
            out.pages.append(
                Page(number=i, text=text, needs_ocr=len(text) < MIN_TEXT_CHARS)
            )

    # Tables are a separate pass — pdfplumber is slower but far better at them.
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                for raw in page.extract_tables() or []:
                    if raw and len(raw) > 1:  # header + at least one row
                        out.tables.append(Table(page=i, rows=raw))
    except Exception:
        # Table extraction is best-effort; never fail ingestion over it.
        pass

    return out


def _parse_docx(data: bytes) -> Parsed:
    doc = DocxDocument(io.BytesIO(data))
    text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    out = Parsed(pages=[Page(number=1, text=text.strip())])
    for t in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in t.rows]
        if len(rows) > 1:
            out.tables.append(Table(page=1, rows=rows))
    return out
